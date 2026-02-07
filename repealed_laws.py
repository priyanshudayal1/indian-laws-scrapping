"""
Scraper for Indian Repealed Laws from https://www.indiacode.nic.in/repealed-act/repealed-act.jsp
Extracts names of all repealed acts and saves them to JSON and Word files.
"""

import argparse
import json
import time
from pathlib import Path

from docx import Document
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait, Select
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException, NoSuchElementException


def scrape_repealed_laws():
    """Scrape all repealed law names from India Code website."""
    
    url = "https://www.indiacode.nic.in/repealed-act/repealed-act.jsp"
    
    # Setup Chrome driver with options
    options = webdriver.ChromeOptions()
    # options.add_argument('--headless')  # Run in headless mode
    options.add_argument('--disable-gpu')
    options.add_argument('--no-sandbox')
    options.add_argument('--disable-dev-shm-usage')
    
    driver = webdriver.Chrome(options=options)
    wait = WebDriverWait(driver, 20)
    
    all_laws = []
    
    try:
        print("Opening webpage...")
        driver.get(url)
        
        # Wait for the table to load
        wait.until(EC.presence_of_element_located((By.ID, "repealedactid")))
        time.sleep(2)  # Allow DataTable to fully initialize
        
        # Change entries per page to 100 for faster scraping
        print("Setting entries per page to 100...")
        try:
            length_select = Select(driver.find_element(By.NAME, "repealedactid_length"))
            length_select.select_by_value("100")
            time.sleep(2)  # Wait for table to reload
        except Exception as e:
            print(f"Could not change page size: {e}")
        
        page_num = 1
        
        while True:
            print(f"Scraping page {page_num}...")
            
            # Wait for table body to be present
            wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "#repealedactid tbody tr")))
            time.sleep(1)
            
            # Extract law names from current page
            rows = driver.find_elements(By.CSS_SELECTOR, "#repealedactid tbody tr")
            
            for row in rows:
                try:
                    cells = row.find_elements(By.TAG_NAME, "td")
                    if len(cells) >= 2:
                        sl_no = cells[0].text.strip()
                        name = cells[1].text.strip()
                        year = cells[2].text.strip() if len(cells) >= 3 else ""
                        
                        if name:
                            all_laws.append({
                                "sl_no": sl_no,
                                "name": name,
                                "year": year
                            })
                except Exception as e:
                    print(f"Error extracting row: {e}")
                    continue
            
            print(f"  Extracted {len(rows)} laws from page {page_num}. Total: {len(all_laws)}")
            
            # Check if there's a next page
            try:
                next_button = driver.find_element(By.ID, "repealedactid_next")
                
                # Check if next button is disabled
                if "disabled" in next_button.get_attribute("class"):
                    print("Reached last page.")
                    break
                
                # Click next page
                next_button.click()
                time.sleep(1.5)  # Wait for page to load
                page_num += 1
                
            except NoSuchElementException:
                print("No more pages.")
                break
            except Exception as e:
                print(f"Error navigating to next page: {e}")
                break
    
    except TimeoutException:
        print("Timeout waiting for page to load.")
    except Exception as e:
        print(f"Error during scraping: {e}")
    finally:
        driver.quit()
    
    return all_laws


def save_to_json(laws, filename="repealed_laws.json"):
    """Save the scraped laws to a JSON file."""
    
    output = {
        "total_count": len(laws),
        "source": "https://www.indiacode.nic.in/repealed-act/repealed-act.jsp",
        "scraped_at": time.strftime("%Y-%m-%d %H:%M:%S"),
        "laws": laws
    }
    
    with open(filename, 'w', encoding='utf-8') as f:
        json.dump(output, f, indent=2, ensure_ascii=False)
    
    print(f"\nSaved {len(laws)} laws to {filename}")


def extract_names_only(laws):
    """Extract only the names of the laws."""
    return [law["name"] for law in laws]


def save_names_docx(names, total, filename="repealed_law_names.docx"):
    """Write the names list to a Word document."""
    doc = Document()
    doc.add_heading("Indian Repealed Laws - Names Only", level=1)
    doc.add_paragraph(f"Total laws: {total}")
    for index, name in enumerate(names, start=1):
        doc.add_paragraph(f"{index}. {name}")
    doc.save(filename)
    print(f"Saved names list to {filename}")


def save_full_docx(laws, filename="repealed_laws.docx"):
    """Write the detailed laws list to a Word document."""
    doc = Document()
    doc.add_heading("Indian Repealed Laws - Detailed Listing", level=1)
    doc.add_paragraph(f"Total laws: {len(laws)}")
    table = doc.add_table(rows=1, cols=3)
    header = table.rows[0].cells
    header[0].text = "Sl. No."
    header[1].text = "Name"
    header[2].text = "Year"
    for entry in laws:
        row = table.add_row().cells
        row[0].text = entry.get("sl_no", "")
        row[1].text = entry.get("name", "")
        row[2].text = entry.get("year", "")
    table.style = "Table Grid"
    doc.save(filename)
    print(f"Saved detailed list to {filename}")


def save_docx_files(laws, names=None, total=None, names_doc="repealed_law_names.docx", detailed_doc="repealed_laws.docx"):
    """Persist Word documents for the given law data."""
    if not laws:
        print("No laws supplied for Word export.")
        return
    names = names if names is not None else extract_names_only(laws)
    total = total if total is not None else len(laws)
    save_full_docx(laws, detailed_doc)
    save_names_docx(names, total, names_doc)


def convert_json_to_word(names_json="repealed_law_names.json", detailed_json="repealed_laws.json", names_doc="repealed_law_names.docx", detailed_doc="repealed_laws.docx"):
    """Convert existing JSON exports into Word documents."""
    generated = False
    detailed_path = Path(detailed_json)
    if detailed_path.exists():
        with detailed_path.open("r", encoding="utf-8") as handle:
            detailed_data = json.load(handle)
        laws = detailed_data.get("laws", [])
        save_full_docx(laws, detailed_doc)
        generated = True
    else:
        print(f"Detailed JSON not found at {detailed_json}")
    names_path = Path(names_json)
    if names_path.exists():
        with names_path.open("r", encoding="utf-8") as handle:
            names_data = json.load(handle)
        names = names_data.get("names", [])
        total = names_data.get("total_count", len(names))
        save_names_docx(names, total, names_doc)
        generated = True
    else:
        print(f"Names JSON not found at {names_json}")
    if not generated:
        print("No Word documents were created.")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Scrape repealed laws or convert existing JSON exports to Word files.")
    parser.add_argument("--convert-only", action="store_true", help="Convert existing JSON files to Word documents without scraping.")
    parser.add_argument("--skip-docx", action="store_true", help="Skip Word document generation after scraping.")
    args = parser.parse_args()

    if args.convert_only:
        convert_json_to_word()
    else:
        print("=" * 60)
        print("Indian Repealed Laws Scraper")
        print("=" * 60)
        laws = scrape_repealed_laws()
        if laws:
            save_to_json(laws, "repealed_laws.json")
            names_only = {
                "total_count": len(laws),
                "names": extract_names_only(laws)
            }
            with open("repealed_law_names.json", "w", encoding="utf-8") as handle:
                json.dump(names_only, handle, indent=2, ensure_ascii=False)
            print("Saved names only to repealed_law_names.json")
            if not args.skip_docx:
                save_docx_files(laws, names_only["names"], names_only["total_count"])
            print("\nFirst 10 repealed laws:")
            print("-" * 40)
            for law in laws[:10]:
                print(f"{law['sl_no']}. {law['name']} ({law['year']})")
        else:
            print("No laws were scraped.")
